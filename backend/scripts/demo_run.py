import json
import os
from pathlib import Path
from docx import Document as DocxDocument
from app.models import ComplianceRun, Document, DocumentChunk, Tenant
from app.services.compliance_service import evaluate_requirements
from app.services.document_service import extract_paragraphs
from app.services.report_service import generate_report
from app.utils.database import Base, SessionLocal, engine


def seed_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("FFIEC 051 Sample Data", level=1)
    doc.add_paragraph("Total assets for Q4 were $1,200,000,000 and total liabilities were $980,000,000.")
    doc.add_paragraph("Tier 1 capital was $120,000,000 and total risk-based capital was $150,000,000.")
    doc.add_paragraph("Past due loans totaled $2,100,000 and nonaccrual loans totaled $700,000.")
    doc.add_paragraph("Loans and leases by category (RC-C) include commercial, consumer, and real estate loans.")
    doc.add_paragraph("Off-balance-sheet items and derivatives (RC-L) include interest rate swaps and loan commitments.")
    doc.save(path)


def main() -> None:
    storage_dir = Path(os.getenv("STORAGE_DIR", "/workspace/Projecthub/backend/app/storage"))
    storage_dir.mkdir(parents=True, exist_ok=True)
    sample_doc = storage_dir / "ffiec051-sample.docx"
    seed_docx(sample_doc)

    Base.metadata.create_all(bind=engine)
    db = SessionLocal()

    tenant = Tenant(name="Demo Bank")
    db.add(tenant)
    db.commit()
    db.refresh(tenant)

    document = Document(
        tenant_id=tenant.id,
        source_type="internal",
        title="FFIEC 051 Sample Bank Data",
        storage_path=str(sample_doc),
    )
    db.add(document)
    db.flush()

    for page_number, paragraph_index, content in extract_paragraphs(sample_doc, "docx"):
        db.add(
            DocumentChunk(
                document_id=document.id,
                page_number=page_number,
                paragraph_index=paragraph_index,
                content=content,
            )
        )
    db.commit()

    requirements_path = Path("/workspace/Projecthub/backend/sample_data/ffiec051_requirements.json")
    requirements = json.loads(requirements_path.read_text())

    run = ComplianceRun(tenant_id=tenant.id, report_type="FFIEC051", status="running")
    db.add(run)
    db.commit()
    db.refresh(run)

    evaluate_requirements(db, run.id, requirements)
    run.status = "completed"
    db.commit()

    report = generate_report(db, run.id, "ffiec051-demo", "pdf")
    print(f"Report generated at: {report.output_path}")

    db.close()


if __name__ == "__main__":
    main()
